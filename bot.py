import os
import time
import sqlite3
import threading
import base64
import json
import io
import traceback
import requests
import telebot

from google import genai
from docx import Document as DocxDocument

from telebot.types import (
    ReplyKeyboardMarkup,
    KeyboardButton,
    InlineKeyboardMarkup,
    InlineKeyboardButton
)


TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "").strip()
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "").strip()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
ADMIN_ID_RAW = os.environ.get("ADMIN_ID", "0").strip()
try:
    ADMIN_ID = int(ADMIN_ID_RAW) if ADMIN_ID_RAW else 0
except ValueError:
    ADMIN_ID = 0

FREE_LIMIT = 15
MONTHLY_PRICE = 100

DB_FILE = "bossai.db"

bot = telebot.TeleBot(TOKEN, parse_mode=None)


def get_db():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_database():
    conn = get_db()
    conn.execute("CREATE TABLE IF NOT EXISTS users (user_id INTEGER PRIMARY KEY, first_name TEXT, username TEXT, free_used INTEGER DEFAULT 0, free_date TEXT, model TEXT DEFAULT 'GPT-4o', subscription_until INTEGER DEFAULT 0, referred_by INTEGER DEFAULT NULL, referrals INTEGER DEFAULT 0, paid_referrals INTEGER DEFAULT 0, created_at INTEGER, notes TEXT DEFAULT '')")
    conn.execute("CREATE TABLE IF NOT EXISTS messages (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, role TEXT, content TEXT, created_at INTEGER)")
    conn.execute("CREATE TABLE IF NOT EXISTS payments (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, amount INTEGER, status TEXT DEFAULT 'pending', created_at INTEGER)")
    conn.execute("CREATE TABLE IF NOT EXISTS feedback (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, rating TEXT, created_at INTEGER)")

    # Safe migration: add notes column if this is an older database file.
    existing_columns = [row["name"] for row in conn.execute("PRAGMA table_info(users)").fetchall()]
    if "notes" not in existing_columns:
        conn.execute("ALTER TABLE users ADD COLUMN notes TEXT DEFAULT ''")
    conn.commit()
    conn.close()


def current_date():
    return time.strftime("%Y-%m-%d")


def get_user(user_id, first_name="", username=""):
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE user_id=?", (user_id,)).fetchone()

    if user is None:
        conn.execute(
            "INSERT INTO users (user_id, first_name, username, free_used, free_date, created_at) VALUES (?, ?, ?, 0, ?, ?)",
            (user_id, first_name or "", username or "", current_date(), int(time.time()))
        )
        conn.commit()
        user = conn.execute("SELECT * FROM users WHERE user_id=?", (user_id,)).fetchone()
    elif user["free_date"] != current_date():
        conn.execute("UPDATE users SET free_used=0, free_date=? WHERE user_id=?", (current_date(), user_id))
        conn.commit()
        user = conn.execute("SELECT * FROM users WHERE user_id=?", (user_id,)).fetchone()

    conn.close()
    return user


def subscription_active(user):
    return user["subscription_until"] and user["subscription_until"] > int(time.time())


def get_subscription_price(user):
    if user["referrals"] >= 50 and user["paid_referrals"] >= 10:
        return 50
    if user["referrals"] >= 30:
        return 70
    return 100


def main_keyboard(user_id=None):
    markup = ReplyKeyboardMarkup(resize_keyboard=True)
    markup.row(KeyboardButton("💳 Payment Methods"), KeyboardButton("👥 Referral"))
    markup.row(KeyboardButton("🤖 Models"), KeyboardButton("🔄 Restart"))
    markup.row(KeyboardButton("❓ Help"), KeyboardButton("📊 My Account"))
    markup.row(KeyboardButton("🎨 Create Image"), KeyboardButton("🎵 Create Music"))
    markup.row(KeyboardButton("📄 Create Document"), KeyboardButton("🧠 My Memory"))
    if user_id is not None and ADMIN_ID != 0 and user_id == ADMIN_ID:
        markup.row(KeyboardButton("👑 Admin Panel"))
    return markup


def save_message(user_id, role, content):
    conn = get_db()
    conn.execute(
        "INSERT INTO messages (user_id, role, content, created_at) VALUES (?, ?, ?, ?)",
        (user_id, role, content, int(time.time()))
    )
    conn.commit()
    conn.close()


def get_history(user_id):
    conn = get_db()
    rows = conn.execute(
        "SELECT role, content FROM messages WHERE user_id=? ORDER BY id DESC LIMIT 10",
        (user_id,)
    ).fetchall()
    conn.close()
    rows = list(reversed(rows))
    return [{"role": row["role"], "content": row["content"]} for row in rows]


def system_prompt(notes=""):
    base = (
        "You are BOSSAI, a natural all-in-one AI assistant. "
        "Default language is English unless the user writes in another language.\n\n"

        "AMHARIC QUALITY (very important): When the user writes in Amharic, you must "
        "respond like a fluent native Amharic speaker, not like a machine translation. "
        "Use natural everyday word order, correct verb conjugation, correct use of "
        "particles (ን, ም, ው, የ, ላይ, ጋር), and idiomatic phrasing an educated native "
        "speaker would actually use in casual conversation. Avoid stiff, overly literal, "
        "or repetitive sentence patterns. Vary sentence length. Do not mix in English "
        "words unless the user did, or unless there is truly no natural Amharic term "
        "(e.g. technical software terms can stay in English). Read your own Amharic "
        "sentence back mentally before answering and make sure it sounds like something "
        "a real person would say, not a direct translation from English.\n\n"

        "If the user writes in another language, respond naturally and fluently in "
        "that same language using the same care described above.\n\n"

        "Do not unnecessarily say that you are a bot. Be helpful, clear and natural. "
        "Remember relevant conversation context.\n\n"

        "IMPORTANT: You cannot generate images, videos, or music yourself through this chat. "
        "If the user asks you to create, draw, generate, or make an image/photo/picture, "
        "do NOT pretend to create one or describe a fake result. Instead, tell them to tap "
        "the '🎨 Create Image' button in the menu to actually generate a real image. "
        "If the user asks you to create, compose, or make music or a song, do NOT pretend "
        "to create one. Instead, tell them to tap the '🎵 Create Music' button in the menu. "
        "Never claim you generated, sent, or attached an image, video, or audio file "
        "unless a real file was actually sent through the system.\n\n"

        "FORMATTING: Write in clean, natural, conversational prose for any language. "
        "Do not use markdown symbols such as **, ##, or excessive bullet dashes. "
        "Use plain sentences and, if a list is genuinely needed, simple short lines "
        "without decorative symbols."
    )

    if notes:
        base += (
            "\n\nWhat you remember about this specific user (use naturally when relevant, "
            "do not just recite it back):\n" + notes
        )

    return base


CHAT_MODELS = {
    "DeepSeek": "deepseek/deepseek-chat",
    "GPT-4o": "openai/gpt-4o",
    "Claude": "anthropic/claude-3.5-sonnet",
    "Grok": "x-ai/grok-2-1212",
}


def ask_openrouter(user_id, text):
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OPENROUTER_API_KEY is missing.")

    user = get_user(user_id)
    model = user["model"]
    history = get_history(user_id)

    messages = [{"role": "system", "content": system_prompt(user["notes"] or "")}]
    messages.extend(history)
    messages.append({"role": "user", "content": text})

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        json={"model": CHAT_MODELS.get(model, CHAT_MODELS["GPT-4o"]), "messages": messages},
        timeout=90
    )

    if not response.ok:
        raise RuntimeError(f"OpenRouter {response.status_code}: {response.text[:300]}")

    data = response.json()
    return data["choices"][0]["message"]["content"]


def ask_gemini(user_id, text):
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is missing.")

    user = get_user(user_id)
    history = get_history(user_id)
    conversation = ""
    for item in history:
        conversation += item["role"] + ": " + item["content"] + "\n"

    prompt = system_prompt(user["notes"] or "") + "\n\nPrevious conversation:\n" + conversation + "\n\nCurrent user message:\n" + text

    client = genai.Client(api_key=GEMINI_API_KEY)
    response = client.models.generate_content(model="gemini-3.6-flash", contents=prompt)

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    return response.text


def ask_ai(user_id, text):
    user = get_user(user_id)

    if user["model"] == "Gemini":
        return ask_gemini(user_id, text)

    try:
        return ask_openrouter(user_id, text)
    except Exception as openrouter_error:
        print("OpenRouter failed:", openrouter_error)
        if GEMINI_API_KEY:
            try:
                return ask_gemini(user_id, text)
            except Exception as gemini_error:
                print("Gemini fallback also failed:", gemini_error)
                raise RuntimeError(f"OpenRouter error: {openrouter_error} | Gemini error: {gemini_error}")
        raise


def notify_admin_error(context, user_id, error):
    if ADMIN_ID == 0:
        return
    try:
        bot.send_message(
            ADMIN_ID,
            f"⚠️ BOSSAI Error\n\nContext: {context}\nUser ID: {user_id}\nError: {str(error)[:400]}"
        )
    except Exception as notify_error:
        print("Could not notify admin of error:", notify_error)


def typing_loop(chat_id, stop_event):
    while not stop_event.is_set():
        try:
            bot.send_chat_action(chat_id, "typing")
        except Exception:
            pass
        stop_event.wait(4)


def clean_formatting(text):
    text = text.replace("**", "").replace("###", "").replace("##", "")
    text = text.replace("* ", "- ")
    return text


def send_long_message(message, text, feedback_markup=None):
    if not text:
        text = "Sorry, I could not generate a response."
    text = clean_formatting(text)

    chunks = [text[i:i + 4000] for i in range(0, len(text), 4000)]
    if not chunks:
        return

    # Reply directly to the user's message so the answer is threaded to their question.
    if len(chunks) == 1:
        bot.reply_to(message, chunks[0], reply_markup=feedback_markup)
        return

    bot.reply_to(message, chunks[0])
    for chunk in chunks[1:-1]:
        bot.send_message(message.chat.id, chunk)
    bot.send_message(message.chat.id, chunks[-1], reply_markup=feedback_markup)


def send_welcome(message, extra_note=""):
    name = message.from_user.first_name or "there"
    text = (
        f"Hello {name}! Welcome to BOSSAI — your all-in-one AI assistant.\n\n"
        "Access GPT-4o, Claude, DeepSeek, Grok, and Gemini in one bot.\n\n"
        "I can:\n"
        "- Answer questions\n"
        "- Write and translate text\n"
        "- Write and debug code\n"
        "- Solve math problems\n"
        "- Remember conversations\n"
        "- Create real images and music\n\n"
        f"Free: {FREE_LIMIT} messages per day\n"
        f"Unlimited: {MONTHLY_PRICE} ETB/month\n\n"
        "Use the buttons below."
    )
    bot.send_message(message.chat.id, text + extra_note, reply_markup=main_keyboard(message.from_user.id))


@bot.message_handler(commands=["start"])
def start(message):
    user_id = message.from_user.id

    conn = get_db()
    already_existed = conn.execute("SELECT 1 FROM users WHERE user_id=?", (user_id,)).fetchone() is not None
    conn.close()

    user = get_user(user_id, message.from_user.first_name, message.from_user.username)

    referral_note = ""

    if message.text:
        parts = message.text.split(maxsplit=1)
        if len(parts) == 2 and parts[1].strip().startswith("ref_"):
            try:
                referrer_id = int(parts[1].strip()[4:])
                if referrer_id != user_id and user["referred_by"] is None:
                    conn = get_db()
                    referrer = conn.execute("SELECT user_id FROM users WHERE user_id=?", (referrer_id,)).fetchone()
                    if referrer:
                        conn.execute(
                            "UPDATE users SET referred_by=? WHERE user_id=? AND referred_by IS NULL",
                            (referrer_id, user_id)
                        )
                        conn.execute("UPDATE users SET referrals=referrals+1 WHERE user_id=?", (referrer_id,))
                        conn.commit()
                        referral_note = "\n\nYou joined through a referral link. Welcome!"
                    conn.close()
            except (ValueError, IndexError):
                pass

    send_welcome(message, referral_note)

    if not already_existed:
        bot.send_message(
            message.chat.id,
            "🎓 Quick Tutorial\n\n"
            "💳 Payment Methods — subscribe for unlimited access\n"
            "👥 Referral — invite friends and unlock discounts\n"
            "🤖 Models — pick which AI model answers you\n"
            "🔄 Restart — clear the current conversation\n"
            "❓ Help — see this info again\n"
            "📊 My Account — check your plan and usage\n"
            "🎨 Create Image — generate a real image\n"
            "🎵 Create Music — generate a real short music clip\n"
            "📄 Create Document — generate a Word file (unlimited subscribers)\n"
            "🧠 My Memory — tell me things to remember about you\n\n"
            "You can also just type any question directly, right now."
        )



@bot.message_handler(commands=["help"])
def help_command(message):
    bot.send_message(
        message.chat.id,
        "BOSSAI Help\n\n"
        "Chat: Send your question directly.\n"
        f"Free: {FREE_LIMIT} messages per day.\n"
        f"Unlimited: {MONTHLY_PRICE} ETB/month.\n"
        "Payment Methods: Choose Telebirr, Payoneer or PayPal.\n"
        "Referral: Invite users and receive discounts.\n"
        "Models: Choose your AI model.\n"
        "Create Image: Generate a real image from a description.\n"
        "Create Music: Generate a real short music clip from a description.\n"
        "Restart: Clear your current conversation.\n\n"
        "Support: @Silent_Survivorr"
    )


@bot.message_handler(func=lambda m: m.text == "❓ Help")
def help_button(message):
    help_command(message)


def show_payment_menu(message):
    user = get_user(message.from_user.id)
    price = get_subscription_price(user)

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton(f"💳 Telebirr — {price} ETB/month", callback_data="telebirr"))
    markup.add(InlineKeyboardButton("🌍 Payoneer", callback_data="payoneer"))
    markup.add(InlineKeyboardButton("🅿️ PayPal", callback_data="paypal"))

    bot.send_message(message.chat.id, "Choose your payment method:", reply_markup=markup)


@bot.message_handler(commands=["menu"])
def menu_command(message):
    show_payment_menu(message)


@bot.message_handler(func=lambda m: m.text == "💳 Payment Methods")
def payment_button(message):
    show_payment_menu(message)


telebirr_waiting = set()


@bot.callback_query_handler(func=lambda call: call.data in ["telebirr", "payoneer", "paypal"])
def payment_callback(call):
    bot.answer_callback_query(call.id)

    if call.data == "telebirr":
        user = get_user(call.from_user.id)
        price = get_subscription_price(user)
        telebirr_waiting.add(call.from_user.id)
        bot.send_message(
            call.message.chat.id,
            f"Telebirr Payment\n\n"
            f"Amount: {price} ETB/month\n\n"
            "Receiver: Hussein\n"
            "Telebirr: 0964990206\n\n"
            "After payment, send your payment receipt screenshot here.\n\n"
            "Your subscription will be activated after manual verification."
        )
    elif call.data == "payoneer":
        bot.send_message(call.message.chat.id, "Payoneer: Soon Available.")
    elif call.data == "paypal":
        bot.send_message(call.message.chat.id, "PayPal: Soon Available.")


@bot.message_handler(content_types=["photo"])
def photo_handler(message):
    user_id = message.from_user.id

    if user_id in telebirr_waiting:
        telebirr_waiting.discard(user_id)
        handle_payment_receipt(message)
        return

    handle_vision_photo(message)


def handle_payment_receipt(message):
    if ADMIN_ID == 0:
        bot.reply_to(message, "Receipt received. Admin verification is not configured yet.")
        return

    user = get_user(message.from_user.id, message.from_user.first_name, message.from_user.username)
    price = get_subscription_price(user)

    conn = get_db()
    cursor = conn.execute(
        "INSERT INTO payments (user_id, amount, status, created_at) VALUES (?, ?, 'pending', ?)",
        (message.from_user.id, price, int(time.time()))
    )
    payment_id = cursor.lastrowid
    conn.commit()
    conn.close()

    markup = InlineKeyboardMarkup()
    markup.add(
        InlineKeyboardButton("✅ Approve", callback_data=f"approve:{payment_id}:{message.from_user.id}"),
        InlineKeyboardButton("❌ Reject", callback_data=f"reject:{payment_id}:{message.from_user.id}")
    )

    caption = (
        f"Payment Receipt\n\nPayment ID: {payment_id}\n"
        f"User: {message.from_user.first_name}\n"
        f"Username: @{message.from_user.username or 'none'}\n"
        f"User ID: {message.from_user.id}\nAmount: {price} ETB\nStatus: Pending"
    )

    bot.send_photo(ADMIN_ID, message.photo[-1].file_id, caption=caption, reply_markup=markup)
    bot.reply_to(message, "Your receipt has been sent for verification. Please wait for approval.")


def handle_vision_photo(message):
    user_id = message.from_user.id
    user = get_user(user_id, message.from_user.first_name, message.from_user.username)

    if not subscription_active(user):
        if user["free_used"] >= FREE_LIMIT:
            bot.reply_to(
                message,
                f"You have used all {FREE_LIMIT} free messages for today.\n\n"
                f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
                "Open Payment Methods to continue."
            )
            return
        conn = get_db()
        conn.execute("UPDATE users SET free_used=free_used+1 WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()

    if not OPENROUTER_API_KEY:
        bot.reply_to(message, "Photo understanding is not available right now.")
        return

    question = (message.caption or "What is in this image? Describe it naturally.").strip()

    stop_event = threading.Event()
    typing_thread = threading.Thread(target=typing_loop, args=(message.chat.id, stop_event), daemon=True)
    typing_thread.start()

    try:
        file_info = bot.get_file(message.photo[-1].file_id)
        file_bytes = bot.download_file(file_info.file_path)
        image_b64 = base64.b64encode(file_bytes).decode("utf-8")

        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": CHAT_MODELS["GPT-4o"],
                "messages": [
                    {"role": "system", "content": system_prompt(user["notes"] or "")},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": question},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}},
                        ],
                    },
                ],
            },
            timeout=90
        )

        if not response.ok:
            raise RuntimeError(f"Vision API {response.status_code}: {response.text[:300]}")

        answer = response.json()["choices"][0]["message"]["content"]
        send_long_message(message, answer)
    except Exception as error:
        print("VISION ERROR:", error)
        traceback.print_exc()
        notify_admin_error("Vision", user_id, error)
        bot.reply_to(message, f"Debug info (temporary): {str(error)[:500]}")
    finally:
        stop_event.set()


@bot.callback_query_handler(func=lambda call: call.data.startswith("approve:") or call.data.startswith("reject:"))
def payment_decision(call):
    if call.from_user.id != ADMIN_ID:
        bot.answer_callback_query(call.id, "Not authorized.", show_alert=True)
        return

    bot.answer_callback_query(call.id)
    parts = call.data.split(":")
    action = parts[0]
    payment_id = int(parts[1])
    user_id = int(parts[2])

    conn = get_db()
    payment = conn.execute("SELECT status FROM payments WHERE id=?", (payment_id,)).fetchone()
    if not payment or payment["status"] != "pending":
        conn.close()
        bot.answer_callback_query(call.id, "This payment was already processed.", show_alert=True)
        return
    conn.close()

    if action == "approve":
        conn = get_db()
        current = conn.execute("SELECT subscription_until FROM users WHERE user_id=?", (user_id,)).fetchone()
        current_until = current["subscription_until"] if current else 0
        base_time = max(int(time.time()), current_until or 0)
        until = base_time + 30 * 24 * 60 * 60

        conn.execute("UPDATE payments SET status='approved' WHERE id=?", (payment_id,))
        conn.execute("UPDATE users SET subscription_until=? WHERE user_id=?", (until, user_id))

        referral = conn.execute("SELECT referred_by FROM users WHERE user_id=?", (user_id,)).fetchone()
        if referral and referral["referred_by"]:
            conn.execute(
                "UPDATE users SET paid_referrals = paid_referrals + 1 WHERE user_id=?",
                (referral["referred_by"],)
            )

        conn.commit()
        conn.close()

        bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id, reply_markup=None)
        bot.send_message(
            user_id,
            "Payment approved.\n\nYour unlimited subscription is active for 30 days.\n\nThank you for using BOSSAI.",
            reply_markup=main_keyboard(user_id)
        )
    else:
        conn = get_db()
        conn.execute("UPDATE payments SET status='rejected' WHERE id=?", (payment_id,))
        conn.commit()
        conn.close()

        bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id, reply_markup=None)
        bot.send_message(
            user_id,
            "Your payment receipt was rejected.\n\nPlease send a valid receipt again.\n\nSupport: @Silent_Survivorr",
            reply_markup=main_keyboard(user_id)
        )


@bot.message_handler(func=lambda m: m.text == "👥 Referral")
def referral(message):
    user = get_user(message.from_user.id)
    bot_username = bot.get_me().username
    referral_link = f"https://t.me/{bot_username}?start=ref_{message.from_user.id}"
    price = get_subscription_price(user)

    conn = get_db()
    invited = conn.execute(
        "SELECT first_name, username FROM users WHERE referred_by=? ORDER BY created_at DESC LIMIT 10",
        (message.from_user.id,)
    ).fetchall()
    conn.close()

    text = (
        f"Referral Program\n\nYour referral link:\n{referral_link}\n\n"
        f"30 referrals gives you 70 ETB/month.\n"
        f"50 referrals plus 10 paid referrals gives you 50 ETB/month.\n\n"
        f"Your referrals: {user['referrals']}\n"
        f"Paid referrals: {user['paid_referrals']}\n\n"
        f"Current price: {price} ETB/month"
    )

    if invited:
        text += "\n\nRecent invites:\n"
        for person in invited:
            person_name = person["first_name"] or "User"
            person_username = f"@{person['username']}" if person["username"] else "no username"
            text += f"- {person_name} ({person_username})\n"

    bot.send_message(message.chat.id, text)


@bot.message_handler(func=lambda m: m.text == "🤖 Models")
def models(message):
    markup = InlineKeyboardMarkup()
    for model in CHAT_MODELS:
        markup.add(InlineKeyboardButton(model, callback_data=f"model:{model}"))
    markup.add(InlineKeyboardButton("Gemini", callback_data="model:Gemini"))

    user = get_user(message.from_user.id)
    bot.send_message(message.chat.id, f"Current model: {user['model']}\n\nChoose a model:", reply_markup=markup)


@bot.callback_query_handler(func=lambda call: call.data.startswith("model:"))
def model_callback(call):
    bot.answer_callback_query(call.id)
    model = call.data.split(":", 1)[1]

    if model not in CHAT_MODELS and model != "Gemini":
        return

    conn = get_db()
    conn.execute("UPDATE users SET model=? WHERE user_id=?", (model, call.from_user.id))
    conn.commit()
    conn.close()

    bot.send_message(call.message.chat.id, f"Model changed to {model}.")


memory_waiting = set()


@bot.message_handler(func=lambda m: m.text == "🧠 My Memory")
def memory_button(message):
    user = get_user(message.from_user.id, message.from_user.first_name, message.from_user.username)
    memory_waiting.add(message.from_user.id)

    current = user["notes"] or "Nothing saved yet."
    bot.reply_to(
        message,
        "🧠 My Memory\n\n"
        f"What I currently remember about you:\n{current}\n\n"
        "Send me anything you want me to remember (your name, your work, your preferences, "
        "things you don't want repeated every time). Send \"clear\" to erase everything I remember."
    )


def process_memory_input(message):
    user_id = message.from_user.id
    memory_waiting.discard(user_id)

    text = (message.text or "").strip()
    if not text:
        bot.reply_to(message, "Please send something to remember, or send \"clear\" to erase memory.")
        return

    conn = get_db()

    if text.lower() in ["clear", "አጥፊ", "አጥፋ", "ሰርዝ"]:
        conn.execute("UPDATE users SET notes='' WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()
        bot.reply_to(message, "🧠 Your memory has been cleared.")
        return

    current = get_user(user_id)["notes"] or ""
    updated = (current + "\n- " + text).strip() if current else "- " + text

    # Keep it from growing unbounded.
    if len(updated) > 2000:
        updated = updated[-2000:]

    conn.execute("UPDATE users SET notes=? WHERE user_id=?", (updated, user_id))
    conn.commit()
    conn.close()

    bot.reply_to(message, "🧠 Got it, I'll remember that.")


@bot.message_handler(func=lambda m: m.text == "📊 My Account")
def account(message):
    user = get_user(message.from_user.id)
    remaining = max(0, FREE_LIMIT - user["free_used"])

    if subscription_active(user):
        days = max(1, int((user["subscription_until"] - int(time.time())) / 86400))
        plan = f"Unlimited active\nApproximately {days} days remaining"
    else:
        plan = "Free plan"

    conn = get_db()
    total_users = conn.execute("SELECT COUNT(*) AS count FROM users").fetchone()["count"]
    paid_users = conn.execute(
        "SELECT COUNT(*) AS count FROM users WHERE subscription_until > ?", (int(time.time()),)
    ).fetchone()["count"]
    conn.close()

    bot.send_message(
        message.chat.id,
        f"My Account\n\nPlan: {plan}\n"
        f"Free messages remaining today: {remaining}\n"
        f"Current model: {user['model']}\n"
        f"Referrals: {user['referrals']}\n"
        f"Paid referrals: {user['paid_referrals']}\n\n"
        f"Registered users: {total_users}\n"
        f"Active paid users: {paid_users}"
    )


@bot.message_handler(func=lambda m: m.text == "🔄 Restart")
def restart(message):
    conn = get_db()
    conn.execute("DELETE FROM messages WHERE user_id=?", (message.from_user.id,))
    conn.commit()
    conn.close()
    send_welcome(message, "\n\n🔄 Your conversation memory has been cleared.")


@bot.message_handler(content_types=["document", "voice", "audio"])
def file_handler(message):
    bot.reply_to(message, "I received your file.\n\nFile and voice analysis can be connected to the appropriate processing service.")


IMAGE_MODEL = "google/gemini-2.5-flash-image"


def generate_image(prompt):
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OPENROUTER_API_KEY is missing.")

    response = requests.post(
        "https://openrouter.ai/api/v1/images",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        json={"model": IMAGE_MODEL, "prompt": prompt},
        timeout=120
    )

    if not response.ok:
        raise RuntimeError(f"Image API {response.status_code}: {response.text[:300]}")

    data = response.json()
    items = data.get("data") or []

    if not items:
        raise RuntimeError("No image data returned.")

    b64 = items[0].get("b64_json")

    if not b64:
        raise RuntimeError("No image data returned.")

    return base64.b64decode(b64)


image_waiting = set()


@bot.message_handler(func=lambda m: m.text == "🎨 Create Image")
def image_button(message):
    image_waiting.add(message.from_user.id)
    bot.send_message(
        message.chat.id,
        "Create Image\n\n"
        "Send me a description of the image you want.\n\n"
        "Example:\nA futuristic city at night, cinematic lighting, realistic, highly detailed."
    )


def process_image_prompt(message):
    user_id = message.from_user.id
    image_waiting.discard(user_id)

    prompt = message.text.strip()

    if not prompt:
        bot.send_message(message.chat.id, "Please describe the image you want.")
        return

    user = get_user(user_id, message.from_user.first_name, message.from_user.username)

    if not subscription_active(user):
        if user["free_used"] >= FREE_LIMIT:
            bot.reply_to(
                message,
                f"You have used all {FREE_LIMIT} free messages for today.\n\n"
                f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
                "Open Payment Methods to continue."
            )
            return
        conn = get_db()
        conn.execute("UPDATE users SET free_used=free_used+1 WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()

    stop_event = threading.Event()
    typing_thread = threading.Thread(target=typing_loop, args=(message.chat.id, stop_event), daemon=True)
    typing_thread.start()

    try:
        bot.send_message(message.chat.id, "Creating your image, please wait...")
        image_bytes = generate_image(prompt)
        bot.send_photo(message.chat.id, image_bytes, caption="Generated by BOSSAI")
    except Exception as error:
        print("IMAGE ERROR:", error)
        traceback.print_exc()
        notify_admin_error("Image generation", user_id, error)
        bot.send_message(message.chat.id, f"Debug info (temporary): {str(error)[:500]}")
    finally:
        stop_event.set()


MUSIC_MODEL = "google/lyria-3-clip-preview"


def generate_music(prompt):
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OPENROUTER_API_KEY is missing.")

    # Audio output on OpenRouter is only delivered via a streamed (SSE) response.
    # A normal non-streaming request will not contain any audio data.
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        json={
            "model": MUSIC_MODEL,
            "modalities": ["text", "audio"],
            "audio": {"format": "mp3"},
            "messages": [{"role": "user", "content": prompt}],
            "stream": True,
        },
        timeout=240,
        stream=True
    )

    if not response.ok:
        raise RuntimeError(f"Music API {response.status_code}: {response.text[:300]}")

    audio_chunks = []

    for line in response.iter_lines():
        if not line:
            continue
        decoded = line.decode("utf-8")
        if not decoded.startswith("data: "):
            continue
        payload = decoded[len("data: "):]
        if payload.strip() == "[DONE]":
            break
        chunk = json.loads(payload)
        choices = chunk.get("choices") or []
        if not choices:
            continue
        audio = (choices[0].get("delta") or {}).get("audio") or {}
        if audio.get("data"):
            audio_chunks.append(audio["data"])

    if not audio_chunks:
        raise RuntimeError("No audio data returned from music model.")

    full_audio_b64 = "".join(audio_chunks)
    return base64.b64decode(full_audio_b64)


music_waiting = set()


@bot.message_handler(func=lambda m: m.text == "🎵 Create Music")
def music_button(message):
    music_waiting.add(message.from_user.id)
    bot.send_message(
        message.chat.id,
        "Create Music\n\n"
        "Describe the song or music you want (genre, mood, instruments).\n\n"
        "Example:\nUpbeat Ethiopian-inspired pop song about friendship, happy mood."
    )


def process_music_prompt(message):
    user_id = message.from_user.id
    music_waiting.discard(user_id)

    prompt = message.text.strip()

    if not prompt:
        bot.send_message(message.chat.id, "Please describe the music you want.")
        return

    user = get_user(user_id, message.from_user.first_name, message.from_user.username)

    if not subscription_active(user):
        if user["free_used"] >= FREE_LIMIT:
            bot.reply_to(
                message,
                f"You have used all {FREE_LIMIT} free messages for today.\n\n"
                f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
                "Open Payment Methods to continue."
            )
            return
        conn = get_db()
        conn.execute("UPDATE users SET free_used=free_used+1 WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()

    stop_event = threading.Event()
    typing_thread = threading.Thread(target=typing_loop, args=(message.chat.id, stop_event), daemon=True)
    typing_thread.start()

    try:
        bot.send_message(message.chat.id, "Composing your music, this can take a minute...")
        audio_bytes = generate_music(prompt)
        audio_file = io.BytesIO(audio_bytes)
        audio_file.name = "bossai_music.mp3"
        bot.send_audio(message.chat.id, audio_file, caption="Generated by BOSSAI")
    except Exception as error:
        print("MUSIC ERROR:", error)
        traceback.print_exc()
        notify_admin_error("Music generation", user_id, error)
        bot.send_message(message.chat.id, f"Debug info (temporary): {str(error)[:500]}")
    finally:
        stop_event.set()


def ask_document_content(topic):
    """Generate well-structured document text (separate from chat history)."""
    document_system_prompt = (
        "You write clean, well-structured documents (reports, letters, essays, "
        "articles, plans, etc.) based on the user's request. "
        "Respond in the same language the user wrote in, with natural, fluent, "
        "professional writing (if Amharic, write like an educated native speaker). "
        "Do not use markdown symbols such as **, ##, or bullet dashes made of *. "
        "Structure the document with a clear title on the first line, followed by "
        "well-organized paragraphs or numbered sections as appropriate. "
        "Do not add commentary about being an AI; just produce the document content."
    )

    if OPENROUTER_API_KEY:
        try:
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": CHAT_MODELS["DeepSeek"],
                    "messages": [
                        {"role": "system", "content": document_system_prompt},
                        {"role": "user", "content": topic},
                    ],
                },
                timeout=120
            )
            if response.ok:
                return response.json()["choices"][0]["message"]["content"]
        except Exception as error:
            print("Document OpenRouter failed:", error)

    if GEMINI_API_KEY:
        client = genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(
            model="gemini-3.6-flash",
            contents=document_system_prompt + "\n\nRequest:\n" + topic
        )
        if response.text:
            return response.text

    raise RuntimeError("No AI service is available to write the document right now.")


def build_docx(text):
    lines = [line.strip() for line in text.strip().split("\n") if line.strip()]
    document = DocxDocument()

    if lines:
        document.add_heading(lines[0], level=1)
        remaining = lines[1:]
    else:
        remaining = lines

    for line in remaining:
        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


doc_waiting = set()


@bot.message_handler(func=lambda m: m.text == "📄 Create Document")
def document_button(message):
    user = get_user(message.from_user.id, message.from_user.first_name, message.from_user.username)

    if not subscription_active(user):
        bot.reply_to(
            message,
            "📄 Document creation is available for unlimited subscribers only.\n\n"
            f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
            "Open Payment Methods to subscribe."
        )
        return

    doc_waiting.add(message.from_user.id)
    bot.reply_to(
        message,
        "📄 Create Document\n\n"
        "Tell me what the document should be about (a report, letter, essay, plan, etc.). "
        "You can be as detailed as you like.\n\n"
        "I will write it and send it back as a Word (.docx) file."
    )


def process_document_prompt(message):
    user_id = message.from_user.id
    doc_waiting.discard(user_id)

    topic = (message.text or "").strip()
    if not topic:
        bot.reply_to(message, "Please describe the document you want.")
        return

    user = get_user(user_id, message.from_user.first_name, message.from_user.username)
    if not subscription_active(user):
        bot.reply_to(
            message,
            "📄 Document creation is available for unlimited subscribers only.\n\n"
            f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
            "Open Payment Methods to subscribe."
        )
        return

    stop_event = threading.Event()
    typing_thread = threading.Thread(target=typing_loop, args=(message.chat.id, stop_event), daemon=True)
    typing_thread.start()

    try:
        bot.reply_to(message, "📄 Writing your document, please wait...")
        content = clean_formatting(ask_document_content(topic))
        docx_file = build_docx(content)
        docx_file.name = "bossai_document.docx"
        bot.send_document(message.chat.id, docx_file, caption="📄 Generated by BOSSAI")
    except Exception as error:
        print("DOCUMENT ERROR:", error)
        traceback.print_exc()
        notify_admin_error("Document generation", user_id, error)
        bot.send_message(message.chat.id, f"Debug info (temporary): {str(error)[:500]}")
    finally:
        stop_event.set()


def is_admin(user_id):
    return ADMIN_ID != 0 and user_id == ADMIN_ID


def admin_panel_markup():
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("📊 Statistics", callback_data="adm:stats"))
    markup.add(
        InlineKeyboardButton("⏳ Pending", callback_data="adm:pending"),
        InlineKeyboardButton("✅ Approved", callback_data="adm:approved"),
        InlineKeyboardButton("❌ Rejected", callback_data="adm:rejected"),
    )
    markup.add(InlineKeyboardButton("⌛ Expiring soon (3 days)", callback_data="adm:expiring"))
    return markup


def build_admin_stats():
    now = int(time.time())
    soon = now + (3 * 86400)
    thirty_days_ago = now - (30 * 86400)

    conn = get_db()

    total_users = conn.execute("SELECT COUNT(*) AS c FROM users").fetchone()["c"]
    active_subs = conn.execute("SELECT COUNT(*) AS c FROM users WHERE subscription_until > ?", (now,)).fetchone()["c"]
    monthly_active = conn.execute(
        "SELECT COUNT(DISTINCT user_id) AS c FROM messages WHERE created_at >= ?", (thirty_days_ago,)
    ).fetchone()["c"]
    expiring_soon = conn.execute(
        "SELECT COUNT(*) AS c FROM users WHERE subscription_until > ? AND subscription_until <= ?",
        (now, soon)
    ).fetchone()["c"]
    pending = conn.execute("SELECT COUNT(*) AS c FROM payments WHERE status='pending'").fetchone()["c"]
    approved = conn.execute("SELECT COUNT(*) AS c FROM payments WHERE status='approved'").fetchone()["c"]
    rejected = conn.execute("SELECT COUNT(*) AS c FROM payments WHERE status='rejected'").fetchone()["c"]
    revenue = conn.execute(
        "SELECT COALESCE(SUM(amount), 0) AS c FROM payments WHERE status='approved'"
    ).fetchone()["c"]

    conn.close()

    return (
        "👑 BOSSAI Admin Dashboard\n\n"
        f"👥 Total users: {total_users}\n"
        f"⭐ Active subscribers: {active_subs}\n"
        f"📅 Active in last 30 days: {monthly_active}\n"
        f"⌛ Expiring within 3 days: {expiring_soon}\n\n"
        f"⏳ Pending payment reviews: {pending}\n"
        f"✅ Approved payments (all time): {approved}\n"
        f"❌ Rejected payments (all time): {rejected}\n"
        f"💰 Approved revenue: {revenue} ETB"
    )


def build_payment_list(status, limit=15):
    conn = get_db()
    rows = conn.execute(
        """
        SELECT payments.id, payments.amount, payments.created_at,
               users.first_name, users.username, users.user_id
        FROM payments
        LEFT JOIN users ON users.user_id = payments.user_id
        WHERE payments.status = ?
        ORDER BY payments.id DESC
        LIMIT ?
        """,
        (status, limit)
    ).fetchall()
    conn.close()

    label = {"pending": "⏳ Pending", "approved": "✅ Approved", "rejected": "❌ Rejected"}[status]

    if not rows:
        return f"{label} payments\n\nNo records found."

    text = f"{label} payments (latest {len(rows)})\n\n"
    for row in rows:
        name = row["first_name"] or "Unknown"
        username = f"@{row['username']}" if row["username"] else "no username"
        when = time.strftime("%Y-%m-%d %H:%M", time.localtime(row["created_at"]))
        text += f"• {name} ({username})\n  ID: {row['user_id']} | {row['amount']} ETB | {when}\n\n"

    return text


def build_expiring_list(limit=15):
    now = int(time.time())
    soon = now + (3 * 86400)

    conn = get_db()
    rows = conn.execute(
        """
        SELECT user_id, first_name, username, subscription_until
        FROM users
        WHERE subscription_until > ? AND subscription_until <= ?
        ORDER BY subscription_until ASC
        LIMIT ?
        """,
        (now, soon, limit)
    ).fetchall()
    conn.close()

    if not rows:
        return "⌛ Expiring within 3 days\n\nNo subscribers expiring soon."

    text = f"⌛ Expiring within 3 days ({len(rows)})\n\n"
    for row in rows:
        name = row["first_name"] or "Unknown"
        username = f"@{row['username']}" if row["username"] else "no username"
        when = time.strftime("%Y-%m-%d %H:%M", time.localtime(row["subscription_until"]))
        text += f"• {name} ({username})\n  ID: {row['user_id']} | expires {when}\n\n"

    return text


@bot.message_handler(commands=["admin"])
def admin_command(message):
    if not is_admin(message.from_user.id):
        return
    bot.send_message(message.chat.id, build_admin_stats(), reply_markup=admin_panel_markup())


@bot.message_handler(func=lambda m: m.text == "👑 Admin Panel")
def admin_panel_button(message):
    if not is_admin(message.from_user.id):
        return
    bot.send_message(message.chat.id, build_admin_stats(), reply_markup=admin_panel_markup())


@bot.callback_query_handler(func=lambda call: call.data.startswith("adm:"))
def admin_panel_callback(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "Not authorized.", show_alert=True)
        return

    bot.answer_callback_query(call.id)
    action = call.data.split(":", 1)[1]

    if action == "stats":
        text = build_admin_stats()
    elif action == "pending":
        text = build_payment_list("pending")
    elif action == "approved":
        text = build_payment_list("approved")
    elif action == "rejected":
        text = build_payment_list("rejected")
    elif action == "expiring":
        text = build_expiring_list()
    else:
        return

    bot.send_message(call.message.chat.id, text, reply_markup=admin_panel_markup())


busy_users = set()
busy_lock = threading.Lock()
last_request = {}


@bot.message_handler(content_types=["text"])
def chat(message):
    text = message.text.strip()
    if not text or text.startswith("/"):
        return

    user_id = message.from_user.id

    if user_id in image_waiting:
        process_image_prompt(message)
        return

    if user_id in music_waiting:
        process_music_prompt(message)
        return

    if user_id in doc_waiting:
        process_document_prompt(message)
        return

    if user_id in memory_waiting:
        process_memory_input(message)
        return

    now = time.time()
    previous = last_request.get(user_id, 0)
    if now - previous < 2:
        bot.reply_to(message, "Wait a moment before your next message.")
        return
    last_request[user_id] = now

    user = get_user(user_id, message.from_user.first_name, message.from_user.username)

    if not subscription_active(user):
        if user["free_used"] >= FREE_LIMIT:
            bot.reply_to(
                message,
                f"You have used all {FREE_LIMIT} free messages for today.\n\n"
                f"Unlimited access is {MONTHLY_PRICE} ETB/month.\n\n"
                "Open Payment Methods to continue."
            )
            return

        conn = get_db()
        conn.execute("UPDATE users SET free_used=free_used+1 WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()

    with busy_lock:
        if user_id in busy_users:
            bot.reply_to(message, "Wait a moment, your previous message is still processing.")
            return
        busy_users.add(user_id)

    stop_event = threading.Event()
    typing_thread = threading.Thread(target=typing_loop, args=(message.chat.id, stop_event), daemon=True)
    typing_thread.start()

    try:
        save_message(user_id, "user", text)
        answer = ask_ai(user_id, text)
        save_message(user_id, "assistant", answer)

        conn = get_db()
        cursor = conn.execute(
            "INSERT INTO feedback (user_id, rating, created_at) VALUES (?, NULL, ?)",
            (user_id, int(time.time()))
        )
        feedback_id = cursor.lastrowid
        conn.commit()
        conn.close()

        feedback_markup = InlineKeyboardMarkup()
        feedback_markup.add(
            InlineKeyboardButton("👍", callback_data=f"fb:{feedback_id}:up"),
            InlineKeyboardButton("👎", callback_data=f"fb:{feedback_id}:down"),
        )

        send_long_message(message, answer, feedback_markup)
    except Exception as error:
        print("CHAT ERROR:", error)
        traceback.print_exc()
        notify_admin_error("Chat", user_id, error)
        bot.reply_to(message, f"Debug info (temporary): {str(error)[:500]}")
    finally:
        stop_event.set()
        with busy_lock:
            busy_users.discard(user_id)


@bot.callback_query_handler(func=lambda call: call.data.startswith("fb:"))
def feedback_callback(call):
    bot.answer_callback_query(call.id, "Thanks for the feedback!")
    parts = call.data.split(":")
    feedback_id = int(parts[1])
    rating = parts[2]

    conn = get_db()
    conn.execute("UPDATE feedback SET rating=? WHERE id=?", (rating, feedback_id))
    conn.commit()
    conn.close()

    try:
        bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id, reply_markup=None)
    except Exception:
        pass


def startup_diagnostic():
    if ADMIN_ID == 0:
        return

    def status_line(name, value):
        if not value:
            return f"NOT SET: {name}"
        tail = value[-4:] if len(value) >= 4 else value
        return f"OK: {name} (ends in {tail}, length {len(value)})"

    report = "BOSSAI Startup Diagnostic\n\n"
    report += status_line("TELEGRAM_BOT_TOKEN", TOKEN) + "\n"
    report += status_line("GEMINI_API_KEY", GEMINI_API_KEY) + "\n"
    report += status_line("OPENROUTER_API_KEY", OPENROUTER_API_KEY) + "\n"
    report += f"ADMIN_ID: {ADMIN_ID}\n"

    try:
        bot.send_message(ADMIN_ID, report)
    except Exception as e:
        print("Could not send startup diagnostic:", e)


def notify_admin(text):
    if ADMIN_ID == 0:
        return
    try:
        bot.send_message(ADMIN_ID, text)
    except Exception as e:
        print("Could not notify admin:", e)


def main():
    try:
        init_database()
    except Exception as error:
        print("Database init error:", error)
        notify_admin(f"BOSSAI failed to initialize the database:\n{str(error)[:500]}")

    print("BOSSAI is running...")

    try:
        bot.remove_webhook()
    except Exception as error:
        print("Could not remove webhook:", error)

    startup_diagnostic()

    while True:
        try:
            bot.infinity_polling(skip_pending=True, timeout=30, long_polling_timeout=30)
        except Exception as error:
            print("Polling error:", error)
            traceback.print_exc()
            notify_admin(f"BOSSAI polling stopped with an error and is retrying:\n{str(error)[:500]}")
            time.sleep(5)


if __name__ == "__main__":
    main()
